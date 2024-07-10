# from flask import Flask, request, render_template, jsonify
# from werkzeug.utils import secure_filename
# import os
# import docx
# import PyPDF2
# import difflib
# import re
# import pandas as pd
# from difflib import SequenceMatcher
# import requests

# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = 'uploads'
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# too_good_phrases = [
#     "Congratulations! You've won the lottery!",
#     "You've been selected to receive a cash prize.",
#     "Claim your inheritance now.",
#     "Get rich quick!",
#     "Receive a free gift card worth $1000.",
#     "Act now before it's too late!",
#     "Limited time offer.",
#     "Offer expires soon.",
#     "Immediate action required.",
#     "Your account will be deactivated if you do not respond within 24 hours.",
#     "Exclusive deal just for you.",
#     "You're one of the lucky few.",
#     "Special offer for loyal customers.",
#     "VIP access only.",
#     "Unlock your exclusive reward.",
#     "Work from home and earn $5000 a week!",
#     "No experience needed. Start earning today.",
#     "Get paid to take surveys.",
#     "Make thousands with our proven system.",
#     "Become a mystery shopper and keep what you buy.",
#     "Your account has been compromised.",
#     "Unusual activity detected on your account.",
#     "Verify your account immediately.",
#     "Update your billing information to avoid suspension.",
#     "Your password needs to be reset.",
#     "Confirm your identity.",
#     "Update your payment details.",
#     "Provide your Social Security number for verification.",
#     "We need additional information to process your request.",
#     "Verify your email address to receive your reward.",
#     "Help a child in need.",
#     "Urgent appeal for disaster relief.",
#     "Your donation can save lives.",
#     "Support our cause and get a free gift.",
#     "Make a difference with just one click."
# ]

# sensitive_phrases = [
#     "Verify your account immediately.",
#     "Update your billing information to avoid suspension.",
#     "Your password needs to be reset.",
#     "Confirm your identity.",
#     "Update your payment details.",
#     "Provide your Social Security number for verification.",
#     "We need additional information to process your request.",
#     "Verify your email address to receive your reward."
# ]

# def check_too_good_phrases(content):
#     found_phrases = []
#     reasons = {}

#     for phrase in too_good_phrases:
#         if phrase.lower() in content.lower():
#             found_phrases.append(phrase)
#             reasons[phrase] = "Too good to be true"
#         else:
#             close_matches = difflib.get_close_matches(phrase.lower(), content.lower().split(), cutoff=0.8)
#             if close_matches:
#                 found_phrases.append(phrase)
#                 reasons[phrase] = "Similar to a too good to be true phrase"

#     for phrase in sensitive_phrases:
#         if phrase.lower() in content.lower():
#             found_phrases.append(phrase)
#             reasons[phrase] = "Sensitive information"

#     if found_phrases:
#         return 'phishing', found_phrases, reasons
#     return 'safe', found_phrases, reasons

# def read_text_file(file_path):
#     with open(file_path, 'r', encoding='utf-8') as file:
#         return file.read()

# def read_docx_file(file_path):
#     doc = docx.Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def read_pdf_file(file_path):
#     reader = PyPDF2.PdfFileReader(file_path)
#     text = ""
#     for page_num in range(reader.getNumPages()):
#         page = reader.getPage(page_num)
#         text += page.extractText()
#     return text

# @app.route('/')
# def index():
#     return render_template('index.html')
    

# @app.route('/check_email', methods=['POST'])
# def check_email():
#     email_content = request.form.get('email_content', '')
#     status, found_phrases, reasons = check_too_good_phrases(email_content)
#     if status == 'phishing':
#         return jsonify({'result': 'Phishing', 'found_phrases': found_phrases, 'reasons': reasons})
#     return jsonify({'result': 'Safe', 'reason': 'No suspicious phrases found in the email.'})

# @app.route('/upload', methods=['POST'])
# def upload():
#     if 'file' not in request.files:
#         return jsonify({'result': 'No file part'})

#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({'result': 'No selected file'})

#     if file:
#         filename = secure_filename(file.filename)
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         file.save(file_path)

#         if filename.lower().endswith('.txt'):
#             content = read_text_file(file_path)
#         elif filename.lower().endswith('.docx'):
#             content = read_docx_file(file_path)
#         elif filename.lower().endswith('.pdf'):
#             content = read_pdf_file(file_path)
#         else:
#             return jsonify({'result': 'Unsupported file type'})

#         status, found_phrases, reasons = check_too_good_phrases(content)
#         if status == 'phishing':
#             return jsonify({'result': 'Phishing', 'found_phrases': found_phrases, 'reasons': reasons})
#         return jsonify({'result': 'Safe', 'reason': 'No suspicious phrases found in the document.'})
#     # Load legitimate domains from the provided CSV file
# csv_path = 'top10milliondomains.csv'
# legitimate_domains = []

# try:
#     df = pd.read_csv(csv_path)
#     if 'Domain' not in df.columns:
#         raise KeyError(f"The 'Domain' column is missing in the CSV file. Available columns: {df.columns.tolist()}")
#     legitimate_domains = df['Domain'].tolist()
# except FileNotFoundError:
#     raise FileNotFoundError("The CSV file was not found. Please make sure 'top10milliondomains.csv' is in the correct directory.")
# except KeyError as e:
#     raise KeyError(e)
# except Exception as e:
#     raise Exception("An error occurred while loading the CSV file: " + str(e))

# # List of known country abbreviations (ccTLDs)
# country_codes = [
#     'us', 'uk', 'za', 'ng', 'ca', 'au', 'de', 'fr', 'it', 'es', 'jp', 'cn', 'in', 'br', 'ru'
# ]

# # Create a pattern for these country codes
# ccTLD_pattern = r'\.(' + '|'.join(country_codes) + ')$'

# # Define allowed URL structures with patterns
# ALLOWED_STRUCTURES = [
#     r'^https://www\.[a-zA-Z0-9-]+\.com/?$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/products/widget$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/category/electronics$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/blog/how-to-use-widgets$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/search\?q=widgets$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/users/johndoe$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/login$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/signup$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/help/contact$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/checkout$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com.ng$',
#     r'^https://www\.[a-zA-Z0-9-]+\.com/[a-zA-Z0-9-]+$',
#     rf'^https://www\.[a-zA-Z0-9-]+\.(com|edu){ccTLD_pattern}',  # New pattern for generic and academic domains with specific ccTLDs
# ]

# # Function to check if the URL matches allowed structures
# def is_allowed_url(url):
#     for pattern in ALLOWED_STRUCTURES:
#         if re.match(pattern, url):
#             return True
#     return False

# # Function to normalize domains by removing prefixes and subdomains
# def normalize_domain(domain):
#     return re.sub(r'^(www\.)?|.*?\.([^.]+\.[^.]+)$', r'\2', domain)

# # Function to extract domain from URL
# def extract_domain(url):
#     match = re.search(r'^(https?://)?([^/]+)', url)
#     if match:
#         return match.group(2)
#     return None

# # Function to perform further domain checks if not found in the top domains list
# def further_domain_checks(domain):
#     try:
#         response = requests.get(f'http://{domain}', timeout=5)
#         if response.status_code == 200:
#             return True
#     except requests.RequestException:
#         return False
#     return False

# # Perform phishing checks
# def check_phishing(url):
#     reasons = []
#     detailed_explanation = []

#     # Check for http instead of https
#     if url.startswith('http://') and not url.startswith('https://'):
#         reasons.append("URL is using HTTP instead of HTTPS, which is less secure.")
#         detailed_explanation.append("Using HTTP instead of HTTPS means the data between your browser and the website is not encrypted. This makes it easier for attackers to intercept and manipulate the data.")

#     # Extract domain from URL
#     domain = extract_domain(url)
#     if not domain:
#         reasons.append("Invalid URL format.")
#         detailed_explanation.append("The URL format seems incorrect, which might indicate it's a phishing attempt.")
#         return reasons, detailed_explanation

#     # Normalize the domain for comparison
#     normalized_domain = normalize_domain(domain)

#     # Check if the domain is in the top 10 million domains list
#     if normalized_domain in map(normalize_domain, legitimate_domains):
#         detailed_explanation.append(f"The domain '{domain}' is found in the top 10 million domains, indicating it is legitimate.")
#         return reasons, detailed_explanation  # Empty reasons, indicating the URL is legitimate

#     # Further checks for misspelled or altered domain names
#     for legit_domain in legitimate_domains:
#         normalized_legit_domain = normalize_domain(legit_domain)
#         if SequenceMatcher(None, normalized_domain, normalized_legit_domain).ratio() > 0.8:
#             if normalized_domain != normalized_legit_domain:
#                 reasons.append(f"URL domain '{domain}' is similar to a legitimate domain '{legit_domain}' but may be misspelled or altered.")
#                 detailed_explanation.append(f"The domain '{domain}' is very similar to '{legit_domain}', which could be an attempt to deceive users by using a look-alike domain.")
#                 break
#     else:
#         # Further domain checks if not found in the top domains list
#         if not further_domain_checks(domain):
#             reasons.append(f"URL domain '{domain}' could not be verified as legitimate.")
#             detailed_explanation.append(f"The domain '{domain}' is not in the top 10 million domains and could not be verified by accessing it, indicating potential risk.")

#     # Check for IP address
#     if re.match(r'^\d+\.\d+\.\d+\.\d+', domain):
#         reasons.append("URL uses an IP address instead of a domain name.")
#         detailed_explanation.append("Legitimate websites usually use domain names, not raw IP addresses. Using an IP address might indicate an attempt to hide the true identity of the site.")

#     # Check for unusual characters
#     if re.search(r'[%\-_\+]', url):
#         reasons.append("URL contains unusual characters like %, -, _, or +.")
#         detailed_explanation.append("The presence of unusual characters in the URL might indicate obfuscation attempts by attackers.")

#     # Check for mismatched URL and content (simplified check here)
#     if "secure-banking" in url and "login" not in url:
#         reasons.append("URL suggests secure banking but does not relate to banking content.")
#         detailed_explanation.append("The URL contains 'secure-banking' but doesn't seem to be related to banking activities, which is suspicious.")

#     # Check for excessive redirects (simplified check here)
#     if "redirect" in url:
#         reasons.append("URL may use excessive redirects to obscure the final destination.")
#         detailed_explanation.append("Excessive redirects can be used to confuse users and hide the final destination of the URL, often used in phishing attempts.")

#     # Check for abnormal path structure
#     if re.search(r'/[a-z0-9]{10,}/', url):
#         reasons.append("URL contains an abnormal path structure with random strings.")
#         detailed_explanation.append("Random strings in the URL path can be an indicator of a phishing attempt trying to create a unique and difficult-to-trace URL.")

#     # Check for file extensions in URL
#     if re.search(r'\.email$', url):
#         reasons.append("URL contains unusual file extension.")
#         detailed_explanation.append("Unusual file extensions in the URL are uncommon for legitimate sites and can indicate a phishing attempt.")

#     return reasons, detailed_explanation



# @app.route('/check_url', methods=['POST'])
# def check_url():
#     data = request.get_json()
#     url = data.get('url', '')

#     # Extract domain from URL
#     domain = extract_domain(url)
#     if not domain:
#         return jsonify({"status": "invalid", "reasons": ["Invalid URL format."], "explanation": ["The URL format seems incorrect, which might indicate it's a phishing attempt."]}), 400

#     # Normalize the domain for comparison
#     normalized_domain = normalize_domain(domain)

#     # Check if the domain is in the top 10 million domains list
#     if normalized_domain in map(normalize_domain, legitimate_domains):
#         return jsonify({"status": "legitimate", "reasons": [], "explanation": [f"The domain '{domain}' is found in the top 10 million domains, indicating it is legitimate."]}), 200

#     # Initial allowed structure check
#     if is_allowed_url(url):
#         reasons, detailed_explanation = check_phishing(url)
#         if reasons:
#             return jsonify({"status": "phishing", "reasons": reasons, "explanation": detailed_explanation}), 400
#         else:
#             return jsonify({"status": "legitimate", "reasons": [], "explanation": ["The URL matches an allowed structure."]}), 200
#     else:
#         # Check for closely related domains if not found in allowed structures
#         reasons, detailed_explanation = check_phishing(url)
#         if reasons:
#             return jsonify({"status": "phishing", "reasons": reasons, "explanation": detailed_explanation}), 400

#     return jsonify({"status": "unknown", "reasons": ["URL does not match allowed structures."], "explanation": ["The URL does not match any of the known allowed structures, which might indicate it's not legitimate."]}), 400


# if __name__ == '__main__':
#     app.run(debug=True)




from flask import Flask, request, render_template, jsonify
from werkzeug.utils import secure_filename
import os
import docx
import PyPDF2
import re
import pandas as pd
import requests
import spacy
import random
import difflib

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

nlp = spacy.load('en_core_web_sm')


# Load legitimate domains from the provided CSV file
csv_path = 'top10milliondomains.csv'
legitimate_domains = []

try:
    df = pd.read_csv(csv_path)
    if 'Domain' not in df.columns:
        raise KeyError(f"The 'Domain' column is missing in the CSV file. Available columns: {df.columns.tolist()}")
    legitimate_domains = df['Domain'].tolist()
except FileNotFoundError:
    raise FileNotFoundError("The CSV file was not found. Please make sure 'top10milliondomains.csv' is in the correct directory.")
except KeyError as e:
    raise KeyError(e)
except Exception as e:
    raise Exception("An error occurred while loading the CSV file: " + str(e))

# List of known country abbreviations (ccTLDs)
country_codes = [
    'us', 'uk', 'za', 'ng', 'ca', 'au', 'de', 'fr', 'it', 'es', 'jp', 'cn', 'in', 'br', 'ru'
]

# Create a pattern for these country codes
ccTLD_pattern = r'\.(' + '|'.join(country_codes) + ')$'

# Define allowed URL structures with patterns
ALLOWED_STRUCTURES = [
    r'^https://www\.[a-zA-Z0-9-]+\.com/?$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/products/widget$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/category/electronics$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/blog/how-to-use-widgets$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/search\?q=widgets$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/users/johndoe$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/login$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/signup$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/help/contact$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/checkout$',
    r'^https://www\.[a-zA-Z0-9-]+\.com.ng$',
    r'^https://www\.[a-zA-Z0-9-]+\.com/[a-zA-Z0-9-]+$',
    rf'^https://www\.[a-zA-Z0-9-]+\.(com|edu){ccTLD_pattern}',  # New pattern for generic and academic domains with specific ccTLDs
]

# Function to check if the URL matches allowed structures
def is_allowed_url(url):
    for pattern in ALLOWED_STRUCTURES:
        if re.match(pattern, url):
            return True
    return False

# Function to normalize domains by removing prefixes and subdomains
def normalize_domain(domain):
    return re.sub(r'^(www\.)?|.*?\.([^.]+\.[^.]+)$', r'\2', domain)

# Function to extract domain from URL
def extract_domain(url):
    match = re.search(r'^(https?://)?([^/]+)', url)
    if match:
        return match.group(2)
    return None

# Function to perform further domain checks if not found in the top domains list
def further_domain_checks(domain):
    try:
        response = requests.get(f'http://{domain}', timeout=5)
        if response.status_code == 200:
            return True
    except requests.RequestException:
        return False
    return False

# Function to generate dynamic explanations
def generate_explanation(reason):
    doc = nlp(reason)
    explanation = []

    for sent in doc.sents:
        explanation.append(str(sent))

    return explanation

# Perform phishing checks
def check_phishing(url):
    reasons = []
    detailed_explanation = []

    # Check for http instead of https
    if url.startswith('http://') and not url.startswith('https://'):
        reasons.append("URL is using HTTP instead of HTTPS, which is less secure.")
        detailed_explanation.extend(generate_explanation("Using HTTP instead of HTTPS means the data between your browser and the website is not encrypted. This makes it easier for attackers to intercept and manipulate the data. Example: 'http://example.com' is less secure than 'https://example.com'."))

    # Extract domain from URL
    domain = extract_domain(url)
    if not domain:
        reasons.append("Invalid URL format.")
        detailed_explanation.extend(generate_explanation("The URL format seems incorrect, which might indicate it's a phishing attempt. Example: 'htp://example' is not a valid URL format."))
        return reasons, detailed_explanation

    # Normalize the domain for comparison
    normalized_domain = normalize_domain(domain)

    # Check if the domain is in the top 10 million domains list
    if normalized_domain in map(normalize_domain, legitimate_domains):
        detailed_explanation.extend(generate_explanation(f"The domain '{domain}' is found in the top 10 million domains, indicating it is legitimate."))
        return reasons, detailed_explanation  # Empty reasons, indicating the URL is legitimate

    # Further checks for misspelled or altered domain names
    for legit_domain in legitimate_domains:
        normalized_legit_domain = normalize_domain(legit_domain)
        if SequenceMatcher(None, normalized_domain, normalized_legit_domain).ratio() > 0.8:
            if normalized_domain != normalized_legit_domain:
                reasons.append(f"URL domain '{domain}' is similar to a legitimate domain '{legit_domain}' but may be misspelled or altered.")
                detailed_explanation.extend(generate_explanation(f"The domain '{domain}' is very similar to '{legit_domain}', which could be an attempt to deceive users by using a look-alike domain. Example: 'goggle.com' instead of 'google.com'."))
                break
    else:
        # Further domain checks if not found in the top domains list
        if not further_domain_checks(domain):
            reasons.append(f"URL domain '{domain}' could not be verified as legitimate.")
            detailed_explanation.extend(generate_explanation(f"The domain '{domain}' is not in the top 10 million domains and could not be verified by accessing it, indicating potential risk."))

    # Check for IP address
    if re.match(r'^\d+\.\d+\.\d+\.\d+', domain):
        reasons.append("URL uses an IP address instead of a domain name.")
        detailed_explanation.extend(generate_explanation("Legitimate websites usually use domain names, not raw IP addresses. Using an IP address might indicate an attempt to hide the true identity of the site. Example: 'http://192.168.1.1' instead of 'http://example.com'."))

    # Check for unusual characters
    if re.search(r'[%\-_\+]', url):
        reasons.append("URL contains unusual characters like %, -, _, or +.")
        detailed_explanation.extend(generate_explanation("The presence of unusual characters in the URL might indicate obfuscation attempts by attackers. Example: 'http://example.com/%20secure'."))

    # Check for mismatched URL and content (simplified check here)
    if "secure-banking" in url and "login" not in url:
        reasons.append("URL suggests secure banking but does not relate to banking content.")
        detailed_explanation.extend(generate_explanation("The URL contains 'secure-banking' but doesn't seem to be related to banking activities, which is suspicious. Example: 'http://secure-banking.example.com/info'."))

    # Check for excessive redirects (simplified check here)
    if "redirect" in url:
        reasons.append("URL may use excessive redirects to obscure the final destination.")
        detailed_explanation.extend(generate_explanation("Excessive redirects can be used to confuse users and hide the final destination of the URL, often used in phishing attempts. Example: 'http://example.com/redirect?to=secure'."))

    # Check for abnormal path structure
    if re.search(r'/[a-z0-9]{10,}/', url):
        reasons.append("URL contains an abnormal path structure with random strings.")
        detailed_explanation.extend(generate_explanation("Random strings in the URL path can be an indicator of a phishing attempt trying to create a unique and difficult-to-trace URL. Example: 'http://example.com/abc123xyz/'."))

    # Check for file extensions in URL
    if re.search(r'\.email$', url):
        reasons.append("URL contains unusual file extension.")
        detailed_explanation.extend(generate_explanation("Unusual file extensions in the URL are uncommon for legitimate sites and can indicate a phishing attempt. Example: 'http://example.email'."))

    return reasons, detailed_explanation









too_good_phrases = [
    "Congratulations! You've won the lottery!",
    "You've been selected to receive a cash prize.",
    "Claim your inheritance now.",
    "Get rich quick!",
    "Receive a free gift card worth $1000.",
    "Act now before it's too late!",
    "Limited time offer.",
    "Offer expires soon.",
    "Immediate action required.",
    "Your account will be deactivated if you do not respond within 24 hours.",
    "Exclusive deal just for you.",
    "You're one of the lucky few.",
    "Special offer for loyal customers.",
    "VIP access only.",
    "Unlock your exclusive reward.",
    "Work from home and earn $5000 a week!",
    "No experience needed. Start earning today.",
    "Get paid to take surveys.",
    "Make thousands with our proven system.",
    "Become a mystery shopper and keep what you buy.",
    "Your account has been compromised.",
    "Unusual activity detected on your account.",
    "Verify your account immediately.",
    "Update your billing information to avoid suspension.",
    "Your password needs to be reset.",
    "Confirm your identity.",
    "Update your payment details.",
    "Provide your Social Security number for verification.",
    "We need additional information to process your request.",
    "Verify your email address to receive your reward.",
    "Help a child in need.",
    "Urgent appeal for disaster relief.",
    "Your donation can save lives.",
    "Support our cause and get a free gift.",
    "Make a difference with just one click."
]

sensitive_phrases = [
    "Verify your account immediately.",
    "Update your billing information to avoid suspension.",
    "Your password needs to be reset.",
    "Confirm your identity.",
    "Update your payment details.",
    "Provide your Social Security number for verification.",
    "We need additional information to process your request.",
    "Verify your email address to receive your reward."
]

# Dynamic explanations for each suspicious phrase
phrase_explanations = {
    "too_good_to_be_true": [
        "The phrase '{phrase}' is commonly used in phishing emails because it promises something extraordinary that is unlikely to be true. Such promises are often used to lure victims into clicking links or providing personal information.",
        "Seeing '{phrase}' should raise a red flag. Offers that seem too good to be true usually are. Phishers use such enticing promises to catch your interest and trick you into taking action.",
        "'{phrase}' suggests an unbelievable offer. Always be skeptical of such claims, as they are a classic sign of phishing attempts aimed at exploiting your desires for quick gains."
    ],
    "request_for_sensitive_info": [
        "The phrase '{phrase}' is asking for sensitive information, which is a common phishing tactic. Legitimate organizations rarely request such details via email.",
        "When you see '{phrase}', be cautious. Phishing emails often try to create a sense of urgency to make you provide personal information without thinking.",
        "'{phrase}' is a request for sensitive data, often used by attackers to steal your identity or financial information. Verify the sender's authenticity before responding."
    ]
}

def get_random_explanation(phrase, category):
    if category in phrase_explanations:
        return random.choice(phrase_explanations[category]).format(phrase=phrase)
    return "The phrase '{phrase}' appears suspicious."

def check_too_good_phrases(content):
    doc = nlp(content)
    found_phrases = []
    reasons = {}

    for phrase in too_good_phrases:
        if phrase.lower() in content.lower():
            found_phrases.append(phrase)
            reasons[phrase] = get_random_explanation(phrase, "too_good_to_be_true")

    for phrase in sensitive_phrases:
        if phrase.lower() in content.lower():
            found_phrases.append(phrase)
            reasons[phrase] = get_random_explanation(phrase, "request_for_sensitive_info")

    if found_phrases:
        return 'phishing', found_phrases, reasons
    return 'safe', found_phrases, reasons

def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def read_docx_file(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def read_pdf_file(file_path):
    reader = PyPDF2.PdfFileReader(file_path)
    text = ""
    for page_num in range(reader.getNumPages()):
        page = reader.getPage(page_num)
        text += page.extractText()
    return text

@app.route('/')
def index():
    return render_template('index.html')
@app.route('/video')
def video():
    return render_template('video/index1.html')
@app.route('/feedback')
def feedback():
    return render_template('feedback/index2.html')

@app.route('/check_email', methods=['POST'])
def check_email():
    email_content = request.form.get('email_content', '')
    status, found_phrases, reasons = check_too_good_phrases(email_content)
    if status == 'phishing':
        return jsonify({'result': 'Phishing', 'found_phrases': found_phrases, 'reasons': reasons})
    return jsonify({'result': 'Safe', 'reason': 'No suspicious phrases found in the email.'})

# @app.route('/check_url', methods=['POST'])
# def check_url():
#     url = request.form.get('url', '')
#     if not url:
#         return jsonify({'result': 'Error', 'reason': 'Please provide a URL or domain name.'})

#     reasons, detailed_explanation = check_phishing(url)
#     if reasons:
#         return jsonify({'result': 'Phishing', 'reasons': reasons, 'detailed_explanation': detailed_explanation})
#     return jsonify({'result': 'Safe', 'reason': 'No phishing indicators found in the URL.'})

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({'result': 'No file part'})

    file = request.files['file']
    if file.filename == '':
        return jsonify({'result': 'No selected file'})

    if file:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        if filename.lower().endswith('.txt'):
            content = read_text_file(file_path)
        elif filename.lower().endswith('.docx'):
            content = read_docx_file(file_path)
        elif filename.lower().endswith('.pdf'):
            content = read_pdf_file(file_path)
        else:
            return jsonify({'result': 'Unsupported file type'})

        status, found_phrases, reasons = check_too_good_phrases(content)
        if status == 'phishing':
            return jsonify({'result': 'Phishing', 'found_phrases': found_phrases, 'reasons': reasons})
        return jsonify({'result': 'Safe', 'reason': 'No suspicious phrases found in the document.'})
    


@app.route('/check_url', methods=['POST'])
def check_url():
    data = request.get_json()
    url = data.get('url', '')

    if not url:
        return jsonify({"status": "invalid", "reasons": ["URL not provided."], "explanation": ["Please provide a URL to check."]}), 400

    # Extract domain from URL
    domain = extract_domain(url)
    if not domain:
        return jsonify({"status": "invalid", "reasons": ["Invalid URL format."], "explanation": ["The URL format seems incorrect, which might indicate it's a phishing attempt."]}), 400

    # Normalize the domain for comparison
    normalized_domain = normalize_domain(domain)

    # Check if the domain is in the top 10 million domains list
    if normalized_domain in map(normalize_domain, legitimate_domains):
        return jsonify({"status": "legitimate", "reasons": [], "explanation": [f"The domain '{domain}' is found in the top 10 million domains, indicating it is legitimate."]}), 200

    # Initial allowed structure check
    if is_allowed_url(url):
        reasons, detailed_explanation = check_phishing(url)
        if reasons:
            return jsonify({"status": "phishing", "reasons": reasons, "explanation": detailed_explanation}), 400
        else:
            return jsonify({"status": "legitimate", "reasons": [], "explanation": ["The URL matches an allowed structure."]}), 200
    else:
        # Check for closely related domains if not found in allowed structures
        reasons, detailed_explanation = check_phishing(url)
        if reasons:
            return jsonify({"status": "phishing", "reasons": reasons, "explanation": detailed_explanation}), 400

    return jsonify({"status": "unknown", "reasons": ["URL does not match allowed structures."], "explanation": ["The URL does not match any of the known allowed structures, which might indicate it's not legitimate."]}), 400


if __name__ == '__main__':
    app.run(debug=True)
